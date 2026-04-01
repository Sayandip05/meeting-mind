from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pathlib import Path
from typing import Literal
from app.models.base import get_db
from app.services.ai_agent_service import AIAgentService
from app.api.auth import get_current_user_id
from app.repositories.meeting_repository import MeetingRepository
from app.core.exceptions import MeetingNotFoundError
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from datetime import datetime

router = APIRouter()

BASE_DIR = Path(__file__).resolve().parent.parent.parent
NOTES_DIR = BASE_DIR / "Notes"
NOTES_DIR.mkdir(exist_ok=True)


@router.post("/generate/{meeting_id}")
def generate_highlights(
    meeting_id: int,
    user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Generate highlights for a meeting"""
    
    # Verify meeting belongs to user
    meeting_repo = MeetingRepository(db)
    meeting = meeting_repo.get_by_id(meeting_id)
    
    if not meeting or meeting.user_id != user_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Meeting not found")
    
    if meeting.status != "completed":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Meeting is still being processed"
        )
    
    # Generate highlights
    ai_service = AIAgentService()
    highlights = ai_service.generate_highlights(meeting_id)
    
    # Update meeting record
    highlights_path = NOTES_DIR / f"highlights_{meeting_id}.txt"
    meeting_repo.update_highlights_path(meeting_id, str(highlights_path))
    
    return {"highlights": highlights}


@router.get("/{meeting_id}")
def get_highlights(
    meeting_id: int,
    user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Get highlights for a meeting"""
    
    meeting_repo = MeetingRepository(db)
    meeting = meeting_repo.get_by_id(meeting_id)
    
    if not meeting or meeting.user_id != user_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Meeting not found")
    
    highlights_path = NOTES_DIR / f"highlights_{meeting_id}.txt"
    
    if not highlights_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Highlights not generated yet")
    
    with open(highlights_path, "r", encoding="utf-8") as f:
        highlights = f.read()
    
    return {"highlights": highlights}


@router.get("/download/{meeting_id}")
def download_highlights(
    meeting_id: int,
    format: Literal["pdf", "txt", "docx"] = Query("pdf"),
    user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Download highlights in specified format"""
    
    meeting_repo = MeetingRepository(db)
    meeting = meeting_repo.get_by_id(meeting_id)
    
    if not meeting or meeting.user_id != user_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Meeting not found")
    
    highlights_path = NOTES_DIR / f"highlights_{meeting_id}.txt"
    
    if not highlights_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Highlights not generated yet")
    
    with open(highlights_path, "r", encoding="utf-8") as f:
        text = f.read()
    
    meeting_name = meeting.name.replace(" ", "_")
    download_time = datetime.now().strftime("%d %b %Y, %I:%M %p")
    
    if format == "txt":
        output_path = NOTES_DIR / f"{meeting_name}_{meeting_id}.txt"
        header = f"Meeting: {meeting.name}\nDownloaded: {download_time}\n\n"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(header + text)
        
        return FileResponse(
            output_path,
            media_type="text/plain",
            filename=f"{meeting.name}.txt"
        )
    
    elif format == "pdf":
        output_path = NOTES_DIR / f"{meeting_name}_{meeting_id}.pdf"
        
        doc = SimpleDocTemplate(str(output_path))
        styles = getSampleStyleSheet()
        elements = []
        
        elements.append(Paragraph("Meeting Highlights", styles["Heading1"]))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph(f"Meeting: {meeting.name}", styles["Normal"]))
        elements.append(Paragraph(f"Downloaded: {download_time}", styles["Normal"]))
        elements.append(Spacer(1, 20))
        
        for line in text.split("\n"):
            if line.strip():
                elements.append(Paragraph(line, styles["BodyText"]))
                elements.append(Spacer(1, 8))
        
        doc.build(elements)
        
        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename=f"{meeting.name}.pdf"
        )
    
    elif format == "docx":
        output_path = NOTES_DIR / f"{meeting_name}_{meeting_id}.docx"
        
        document = Document()
        document.add_heading("Meeting Highlights", 0)
        document.add_paragraph(f"Meeting: {meeting.name}")
        document.add_paragraph(f"Downloaded: {download_time}")
        document.add_paragraph("")
        
        for line in text.split("\n"):
            if line.strip():
                document.add_paragraph(line)
        
        document.save(output_path)
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{meeting.name}.docx"
        )
    
    else:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid format")
